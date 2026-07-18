import io
import os
import re
import logging
import json
import zipfile

import pdfplumber
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.infrastructure.vector_store import ChromaDBStore

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    try:
        from pypdf import PdfReader
    except ImportError:
        PdfReader = None

try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)


class UploadUseCase:
    def __init__(self, vector_store: ChromaDBStore):
        self.vector_store = vector_store

    def _ocr_image(self, image_bytes: bytes) -> str:
        if not pytesseract or not Image:
            return ""
        try:
            # Configure tesseract paths for Windows if on Windows and not already done
            if os.name == "nt" and not getattr(self, "_tesseract_configured", False):
                import shutil
                tesseract_candidates = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                    r"C:\msys64\ucrt64\bin\tesseract.exe",
                    r"C:\msys64\mingw64\bin\tesseract.exe",
                ]
                tesseract_path = shutil.which("tesseract")
                if not tesseract_path:
                    for candidate in tesseract_candidates:
                        if os.path.exists(candidate):
                            tesseract_path = candidate
                            break
                if tesseract_path:
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path
                    logger.info(f"Tesseract OCR configured on Windows at: {tesseract_path}")
                self._tesseract_configured = True

            img = Image.open(io.BytesIO(image_bytes))
            # Run OCR supporting Vietnamese and English
            text = pytesseract.image_to_string(img, lang="vie+eng")
            return text.strip()
        except Exception as e:
            logger.warning(f"OCR failed or tesseract not installed: {e}")
            return ""

    def _extract_text_docx_xml(self, file_content: bytes) -> str:
        import xml.etree.ElementTree as ET
        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as z:
                xml_content = z.read("word/document.xml")
                root = ET.fromstring(xml_content)
                texts = []
                for elem in root.iter():
                    if elem.tag.endswith('}t'): # w:t text
                        if elem.text:
                            texts.append(elem.text)
                    elif elem.tag.endswith('}br'): # w:br line break
                        texts.append('\n')
                    elif elem.tag.endswith('}p'): # w:p paragraph
                        texts.append('\n\n')
                return "".join(texts)
        except Exception as e:
            logger.warning(f"XML docx text extraction failed: {e}")
            return ""

    def _extract_text_from_doc(self, file_content: bytes, filename: str) -> str:
        import subprocess
        import tempfile

        # Strategy 1: Try python-docx (works if .doc is actually a .docx saved with wrong extension)
        if docx:
            try:
                doc_obj = docx.Document(io.BytesIO(file_content))
                text_parts = []
                for para in doc_obj.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                for table in doc_obj.tables:
                    for row in table.rows:
                        row_cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                        text_parts.append("| " + " | ".join(row_cells) + " |")
                result_text = "\n\n".join(text_parts)
                if len(result_text.strip()) > 50:
                    logger.info("Successfully extracted .doc text via python-docx (file is actually docx format).")
                    return result_text
            except Exception as e:
                logger.debug(f"python-docx failed for .doc file (expected for true binary .doc): {e}")

        # Strategy 2: Try docx2txt
        try:
            import docx2txt
            text = docx2txt.process(io.BytesIO(file_content))
            if text and len(text.strip()) > 50:
                logger.info("Successfully extracted .doc text via docx2txt.")
                return text
        except Exception as e:
            logger.debug(f"docx2txt failed for .doc: {e}")

        # Strategy 3: Try OLE compound document parsing (true .doc binary format)
        try:
            import olefile
            if olefile.isOleFile(io.BytesIO(file_content)):
                ole = olefile.OleFileIO(io.BytesIO(file_content))
                if ole.exists('WordDocument'):
                    # Try to find the text in the Word Document stream
                    # The actual text in .doc files is stored as UTF-16LE in certain streams
                    text_parts = []
                    for stream_name in ['WordDocument', '1Table', '0Table']:
                        if ole.exists(stream_name):
                            try:
                                stream_data = ole.openstream(stream_name).read()
                                # Try UTF-16LE decoding (common for .doc text)
                                decoded = stream_data.decode('utf-16-le', errors='ignore')
                                # Filter to meaningful text segments
                                import re as re_mod
                                segments = re_mod.findall(r'[\w\s.,;:!?\-\'"()]{10,}', decoded)
                                text_parts.extend(segments)
                            except Exception:
                                pass
                    ole.close()
                    combined = "\n\n".join(text_parts)
                    if len(combined.strip()) > 50:
                        logger.info("Successfully extracted .doc text via OLE compound document parsing.")
                        return combined
                ole.close()
        except ImportError:
            logger.debug("olefile not installed, skipping OLE parsing.")
        except Exception as e:
            logger.debug(f"OLE parsing failed for .doc: {e}")

        # Strategy 4: Try antiword CLI
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_doc:
            temp_doc.write(file_content)
            temp_doc_path = temp_doc.name

        try:
            result = subprocess.run(["antiword", temp_doc_path], capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and len(result.stdout.strip()) > 50:
                logger.info("Successfully extracted text using antiword.")
                return result.stdout
        except Exception as e:
            logger.warning(f"antiword execution failed or not installed: {e}")
        finally:
            try:
                os.remove(temp_doc_path)
            except Exception:
                pass

        # Strategy 5: Extract printable strings (last resort)
        logger.info("Using printable strings fallback for legacy .doc file.")
        import string
        printable_chars = set(string.printable.encode('ascii'))
        text_chunks = []
        current_chunk = []
        for b in file_content:
            if b in printable_chars:
                current_chunk.append(b)
            else:
                if len(current_chunk) >= 8:
                    try:
                        decoded = bytes(current_chunk).decode('utf-8', errors='ignore').strip()
                        if decoded and not decoded.startswith('<?xml') and 'schemas.openxmlformats' not in decoded:
                            text_chunks.append(decoded)
                    except Exception:
                        pass
                current_chunk = []
        if len(current_chunk) >= 8:
            try:
                decoded = bytes(current_chunk).decode('utf-8', errors='ignore').strip()
                if decoded and not decoded.startswith('<?xml') and 'schemas.openxmlformats' not in decoded:
                    text_chunks.append(decoded)
            except Exception:
                pass

        return "\n\n".join(text_chunks)

    def execute(self, file_content: bytes, filename: str, notebook_id: str = "default") -> dict:
        logger.debug(f"File size received: {len(file_content)} bytes")
        
        # Ensure static outputs folder exists
        os.makedirs("app/static/outputs", exist_ok=True)
        
        # We will collect chunks and their metadatas
        all_chunks = []
        all_metadatas = []
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", " ", ""]
        )

        ext = filename.lower()
        if ext.endswith(".pdf"):
            # PDF page-by-page processing
            # 1. Open with pdfplumber for text/tables
            # 2. Open with PyPDF2 (or pypdf) for images
            pdf_reader = None
            processed_images = {}  # Cache to avoid duplicate image processing & OCR
            if PdfReader:
                try:
                    pdf_reader = PdfReader(io.BytesIO(file_content))
                except Exception as e:
                    logger.warning(f"PyPDF2 reader failed to load: {e}")
                
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                logger.debug(f"Total pages parsed: {len(pdf.pages)}")
                for page_idx, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    
                    # Extract tables from this page
                    tables = page.extract_tables() or []
                    for table in tables:
                        if not table:
                            continue
                        for row in table:
                            if row is None:
                                continue
                            clean_row = []
                            for cell in row:
                                if cell is None:
                                    clean_row.append(" ")
                                else:
                                    clean_str = str(cell).replace('\n', ' ').replace('\r', '').strip()
                                    clean_row.append(clean_str if clean_str else " ")
                            markdown_row = "| " + " | ".join(clean_row) + " |\n"
                            page_text += "\n" + markdown_row
                    
                    # Extract images from this page
                    image_urls = []
                    ocr_texts = []
                    if pdf_reader and page_idx < len(pdf_reader.pages):
                        try:
                            reader_page = pdf_reader.pages[page_idx]
                            for img_idx, image_file_object in enumerate(reader_page.images):
                                try:
                                    img_data = image_file_object.data
                                    if not img_data:
                                        continue
                                        
                                    # 1. Filter out tiny decorative images (icons, lines, spacers) below 40x40 pixels
                                    if Image:
                                        try:
                                            with Image.open(io.BytesIO(img_data)) as img_obj:
                                                width, height = img_obj.size
                                                if width < 40 or height < 40:
                                                    logger.debug(f"Skipping tiny decorative image on page {page_idx}: {width}x{height}")
                                                    continue
                                        except Exception as pil_err:
                                            logger.debug(f"Failed to check image size via PIL: {pil_err}")

                                    # 2. Skip duplicate images (like header/footer logos) using MD5 hash
                                    import hashlib
                                    img_hash = hashlib.md5(img_data).hexdigest()
                                    
                                    if img_hash in processed_images:
                                        logger.info(f"Reusing cached image and OCR for duplicate hash {img_hash} on page {page_idx}")
                                        cached_url, cached_ocr = processed_images[img_hash]
                                        image_urls.append(cached_url)
                                        if cached_ocr:
                                            ocr_texts.append(cached_ocr)
                                        continue

                                    # Clean image name
                                    clean_filename = re.sub(r"[^A-Za-z0-9._-]", "_", filename)
                                    img_name = f"{clean_filename}_p{page_idx}_img{img_idx}_{re.sub(r'[^A-Za-z0-9._-]', '_', image_file_object.name)}"
                                    img_path = os.path.join("app/static/outputs", img_name)
                                    
                                    with open(img_path, "wb") as img_f:
                                        img_f.write(img_data)
                                        
                                    image_url = f"/static/outputs/{img_name}"
                                    image_urls.append(image_url)
                                    
                                    # Perform OCR
                                    ocr_text = self._ocr_image(img_data)
                                    if ocr_text:
                                        ocr_texts.append(ocr_text)
                                        
                                    # Cache for deduplication
                                    processed_images[img_hash] = (image_url, ocr_text)
                                except Exception as img_err:
                                    logger.warning(f"Failed to process image {img_idx} on page {page_idx}: {img_err}")
                        except Exception as page_err:
                            logger.warning(f"Failed to extract images from page {page_idx}: {page_err}")
                            
                    # Append OCR text to page content to make it searchable
                    if ocr_texts:
                        page_text += "\n\n[Nội dung hình ảnh trích xuất qua OCR:\n" + "\n".join(ocr_texts) + "\n]"
                        
                    # Split page text into chunks
                    if page_text.strip():
                        chunks = text_splitter.split_text(page_text)
                        for i, chunk in enumerate(chunks):
                            meta = {
                                "source": filename,
                                "chunk_index": len(all_chunks),
                                "notebook_id": notebook_id,
                                "page": page_idx + 1
                            }
                            if image_urls:
                                meta["images"] = json.dumps(image_urls)
                            all_chunks.append(chunk)
                            all_metadatas.append(meta)

        elif ext.endswith(".docx"):
            # docx processing
            doc_text = ""
            image_urls = []
            ocr_texts = []
            
            # Extract text/tables
            if docx:
                try:
                    doc = docx.Document(io.BytesIO(file_content))
                    for para in doc.paragraphs:
                        if para.text.strip():
                            doc_text += para.text + "\n\n"
                    for table in doc.tables:
                        for row in table.rows:
                            row_cells = [cell.text.replace('\n', ' ').replace('\r', '').strip() for cell in row.cells]
                            doc_text += "| " + " | ".join(row_cells) + " |\n"
                        doc_text += "\n"
                except Exception as docx_err:
                    logger.warning(f"python-docx parsing failed: {docx_err}. Using XML fallback.")
                    doc_text = self._extract_text_docx_xml(file_content)
            else:
                doc_text = self._extract_text_docx_xml(file_content)
                
            # Extract images from zip
            try:
                with zipfile.ZipFile(io.BytesIO(file_content)) as z:
                    img_idx = 0
                    for z_name in z.namelist():
                        if z_name.startswith("word/media/"):
                            try:
                                img_data = z.read(z_name)
                                img_ext = os.path.splitext(z_name)[1] or ".png"
                                clean_filename = re.sub(r"[^A-Za-z0-9._-]", "_", filename)
                                img_name = f"{clean_filename}_img{img_idx}{img_ext}"
                                img_path = os.path.join("app/static/outputs", img_name)
                                
                                with open(img_path, "wb") as img_f:
                                    img_f.write(img_data)
                                    
                                image_url = f"/static/outputs/{img_name}"
                                image_urls.append(image_url)
                                
                                ocr_text = self._ocr_image(img_data)
                                if ocr_text:
                                    ocr_texts.append(ocr_text)
                                    
                                img_idx += 1
                            except Exception as e:
                                logger.warning(f"Error extracting docx image {z_name}: {e}")
            except Exception as zip_err:
                logger.warning(f"Failed to extract images from docx zip: {zip_err}")
                
            if ocr_texts:
                doc_text += "\n\n[Nội dung hình ảnh trích xuất qua OCR:\n" + "\n".join(ocr_texts) + "\n]"
                
            if doc_text.strip():
                chunks = text_splitter.split_text(doc_text)
                for i, chunk in enumerate(chunks):
                    meta = {
                        "source": filename,
                        "chunk_index": len(all_chunks),
                        "notebook_id": notebook_id
                    }
                    if image_urls:
                        meta["images"] = json.dumps(image_urls)
                    all_chunks.append(chunk)
                    all_metadatas.append(meta)

        elif ext.endswith(".doc"):
            # doc legacy processing
            doc_text = self._extract_text_from_doc(file_content, filename)
            if doc_text.strip():
                chunks = text_splitter.split_text(doc_text)
                for i, chunk in enumerate(chunks):
                    meta = {
                        "source": filename,
                        "chunk_index": len(all_chunks),
                        "notebook_id": notebook_id
                    }
                    all_chunks.append(chunk)
                    all_metadatas.append(meta)

        elif ext.endswith((".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp")):
            # Save the image to static outputs
            clean_filename = re.sub(r"[^A-Za-z0-9._-]", "_", filename)
            img_name = f"{clean_filename}"
            img_path = os.path.join("app/static/outputs", img_name)
            with open(img_path, "wb") as img_f:
                img_f.write(file_content)
            
            image_url = f"/static/outputs/{img_name}"
            image_urls = [image_url]
            
            # Run OCR on the image
            ocr_text = self._ocr_image(file_content)
            if ocr_text and ocr_text.strip():
                chunks = text_splitter.split_text(ocr_text)
            else:
                # No text found by OCR — store a descriptive placeholder so the image is still indexed
                chunks = [f"[Hình ảnh: {filename}]"]
            
            for i, chunk in enumerate(chunks):
                meta = {
                    "source": filename,
                    "chunk_index": len(all_chunks),
                    "notebook_id": notebook_id
                }
                meta["images"] = json.dumps(image_urls)
                all_chunks.append(chunk)
                all_metadatas.append(meta)

        else:
            # Plain text
            combined_text = file_content.decode("utf-8", errors="ignore")
            if combined_text.strip():
                chunks = text_splitter.split_text(combined_text)
                for i, chunk in enumerate(chunks):
                    meta = {
                        "source": filename,
                        "chunk_index": len(all_chunks),
                        "notebook_id": notebook_id
                    }
                    all_chunks.append(chunk)
                    all_metadatas.append(meta)

        if all_chunks:
            self.vector_store.add_documents(texts=all_chunks, metadatas=all_metadatas)

        logger.info(f"Processed '{filename}': {len(all_chunks)} chunks stored for notebook '{notebook_id}'.")
        return {
            "filename": filename,
            "total_chunks": len(all_chunks),
            "message": f"Successfully processed and stored {len(all_chunks)} chunks."
        }
